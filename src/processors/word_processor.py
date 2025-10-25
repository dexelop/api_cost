"""
Word 문서 프로세서

Word 문서에서 텍스트를 추출하여 처리합니다.
"""

from pathlib import Path

from docx import Document

from .base import BaseFileProcessor, ProcessedFile


class WordFileProcessor(BaseFileProcessor):
    """
    Word 문서 프로세서

    python-docx를 사용하여 Word 문서에서 텍스트를 추출합니다.
    """

    def __init__(self):
        """Word 프로세서 초기화"""
        super().__init__()
        self.supported_extensions = [".docx", ".docm"]

    def extract_paragraphs(self, doc: Document) -> list[str]:
        """
        문서의 모든 단락에서 텍스트 추출

        Args:
            doc: python-docx Document 객체

        Returns:
            list: 각 단락의 텍스트 리스트
        """
        paragraphs = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:  # 빈 단락은 제외
                paragraphs.append(text)
        return paragraphs

    def extract_tables(self, doc: Document) -> list[str]:
        """
        문서의 모든 표에서 데이터 추출

        Args:
            doc: python-docx Document 객체

        Returns:
            list: 각 표의 텍스트 리스트
        """
        tables_text = []

        for table_idx, table in enumerate(doc.tables):
            table_data = []
            table_data.append(f"\n=== Table {table_idx + 1} ===")

            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(" | ".join(row_data))

            tables_text.append("\n".join(table_data))

        return tables_text

    def extract_all_text(self, doc: Document) -> str:
        """
        문서의 모든 텍스트 (단락 + 표) 추출

        Args:
            doc: python-docx Document 객체

        Returns:
            str: 추출된 전체 텍스트
        """
        all_text = []

        # 단락 텍스트 추출
        paragraphs = self.extract_paragraphs(doc)
        all_text.extend(paragraphs)

        # 표 텍스트 추출
        tables = self.extract_tables(doc)
        all_text.extend(tables)

        return "\n\n".join(all_text)

    def get_word_metadata(self, doc: Document, file_path: Path) -> dict:
        """
        Word 문서의 메타데이터 추출

        Args:
            doc: python-docx Document 객체
            file_path: Word 파일 경로

        Returns:
            dict: Word 문서 메타데이터
        """
        # 기본 파일 메타데이터
        metadata = self.get_file_metadata(file_path)

        # Word 전용 메타데이터 추가
        word_metadata = {
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "section_count": len(doc.sections),
        }

        # 문서 속성 (core properties) 추가
        try:
            core_props = doc.core_properties
            word_metadata.update(
                {
                    "title": core_props.title or "",
                    "author": core_props.author or "",
                    "subject": core_props.subject or "",
                    "keywords": core_props.keywords or "",
                    "comments": core_props.comments or "",
                    "created": str(core_props.created) if core_props.created else "",
                    "modified": str(core_props.modified) if core_props.modified else "",
                }
            )
        except Exception:
            pass  # 문서 속성이 없는 경우 무시

        metadata.update(word_metadata)
        return metadata

    def process(self, file_path: Path) -> ProcessedFile:
        """
        Word 문서를 처리하여 텍스트 추출

        Args:
            file_path: 처리할 Word 파일 경로

        Returns:
            ProcessedFile: 처리된 파일 정보
        """
        try:
            # 파일 검증
            self.validate_file(file_path)

            # Word 문서 열기
            doc = Document(str(file_path))

            # 모든 텍스트 추출
            content = self.extract_all_text(doc)

            # 메타데이터 추출
            metadata = self.get_word_metadata(doc, file_path)

            # 텍스트 관련 메타데이터 추가
            metadata.update(
                {
                    "character_count": len(content),
                    "word_count": len(content.split()),
                }
            )

            return ProcessedFile(
                file_path=file_path,
                file_type="word",
                content=content,
                metadata=metadata,
                error=None,
            )

        except Exception as e:
            # 에러 발생 시 에러 정보와 함께 반환
            return ProcessedFile(
                file_path=file_path,
                file_type="word",
                content="",
                metadata=self.get_file_metadata(file_path)
                if file_path.exists()
                else {},
                error=str(e),
            )
